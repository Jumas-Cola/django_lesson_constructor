from django.shortcuts import render
from django.shortcuts import get_object_or_404

from django.contrib.auth.models import User

from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.decorators import login_required
# from django.contrib.auth.mixins import PermissionRequiredMixin
# from django.contrib.auth.decorators import permission_required
# from django.contrib.auth.models import User

from django.http import HttpResponse
from django.http import Http404
# from django.http import HttpResponseRedirect

from django.views import generic
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.views.generic.list import MultipleObjectMixin

from django.urls import reverse_lazy
# from django.urls import reverse

from docx import Document
import os
import json
import uuid
from datetime import date
from .models import TeachingMethod, LessonPart, Comment


def index(request):
    author_name = 'admin'
    parts_with_items = {}
    lesson_parts = LessonPart.objects.all()
    for part in lesson_parts:
        parts_with_items[part] = TeachingMethod.objects.filter(
            author__username__exact=author_name,
            lesson_part__exact=part
        )

    return render(
        request,
        'index.html',
        context={
            'parts_with_items': parts_with_items,
            'lesson_parts': lesson_parts,
        },
    )


# @login_required
def my_methods(request, pk):
    # if request.user.id != int(pk):
    #     raise Http404
    author = get_object_or_404(User, pk=pk)
    parts_with_items = {}
    lesson_parts = LessonPart.objects.all()
    for part in lesson_parts:
        parts_with_items[part] = TeachingMethod.objects.filter(
            author__id__exact=pk,
            lesson_part__exact=part
        )

    return render(
        request,
        'constructor/my_methods.html',
        context={
            'parts_with_items': parts_with_items,
            'lesson_parts': lesson_parts,
            'author': author,
        },
    )


class MethodDetailView(generic.DetailView, MultipleObjectMixin):
    model = TeachingMethod
    template_name = 'constructor/method_detail.html'
    paginate_by = 10

    def get_context_data(self, **kwargs):
        object_list = Comment.objects.filter(method=self.get_object())
        context = super().get_context_data(object_list=object_list, **kwargs)
        return context

    # def get_object(self, queryset=None):
    #     obj = super().get_object()
    #     if not obj.author == self.request.user:
    #         raise Http404
    #     return obj


class MethodDelete(LoginRequiredMixin, DeleteView):
    model = TeachingMethod
    template_name = 'constructor/method_confirm_delete.html'

    def get_success_url(self):
        return self.request.GET['next']

    def get_object(self, queryset=None):
        obj = super().get_object()
        if not obj.author == self.request.user:
            raise Http404
        return obj


class MethodUpdate(LoginRequiredMixin, UpdateView):
    model = TeachingMethod
    template_name = 'constructor/method_form.html'
    fields = ('title', 'description', 'lesson_part')

    def get_success_url(self):
        return self.request.GET['next']

    def get_object(self, queryset=None):
        obj = super().get_object()
        if not obj.author == self.request.user:
            raise Http404
        return obj


class MethodCreate(LoginRequiredMixin, CreateView):
    model = TeachingMethod
    template_name = 'constructor/method_form.html'
    fields = ('title', 'description', 'lesson_part')

    def get_success_url(self):
        return self.request.GET['next']

    def form_valid(self, form):
        form.instance.author = self.request.user
        return super().form_valid(form)


class MethodCopy(LoginRequiredMixin, UpdateView):
    model = TeachingMethod
    template_name = 'constructor/method_form.html'
    fields = ('title', 'description', 'lesson_part')

    def get_success_url(self):
        return self.request.GET['next']

    def form_valid(self, form):
        form.instance.id = uuid.uuid4()
        form.instance.author = self.request.user
        return super().form_valid(form)


def download(request):
    ids_json = request.POST.get('ids')
    ids = json.loads(ids_json)
    document = Document()
    docx_title = "lesson_plan_{}.docx".format(date.today())
    # ---- Cover Letter ----
    document.add_heading('План урока', 0)

    for id in ids:
        method = TeachingMethod.objects.get(pk=id)
        document.add_heading(method.title, level=2)
        document.add_paragraph(method.description)
    document.add_page_break()
    document.save(docx_title)

    file_path = os.path.realpath(docx_title)
    data = open(file_path, "rb").read()
    response = HttpResponse(data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = os.path.getsize(file_path)
    os.remove(file_path)

    return response


class CommentCreate(LoginRequiredMixin, CreateView):
    model = Comment
    fields = ('text',)

    def get_success_url(self):
        return reverse_lazy('method_detail', kwargs={'pk': self.kwargs['pk']})

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['method'] = get_object_or_404(TeachingMethod, pk=self.kwargs['pk'])
        context['method_pk'] = self.kwargs['pk']
        return context

    def form_valid(self, form):
        form.instance.author = self.request.user
        form.instance.method = get_object_or_404(TeachingMethod, pk=self.kwargs['pk'])
        return super().form_valid(form)


class CommentDelete(LoginRequiredMixin, DeleteView):
    model = Comment

    def get_success_url(self):
        return self.request.GET['next']

    def get_object(self, queryset=None):
        obj = super().get_object()
        if not obj.author == self.request.user:
            raise Http404
        return obj


class CommentUpdate(LoginRequiredMixin, UpdateView):
    model = Comment
    fields = ('text',)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['method'] = get_object_or_404(TeachingMethod, pk=self.kwargs['id'])
        context['method_pk'] = self.kwargs['id']
        return context

    def get_success_url(self):
        return self.request.GET['next']

    def get_object(self, queryset=None):
        obj = super().get_object()
        if not obj.author == self.request.user:
            raise Http404
        return obj

    def form_valid(self, form):
        form.instance.is_modified = True
        return super().form_valid(form)
